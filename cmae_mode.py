import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

def carregar_dados(uploaded_file):
    """Carrega os dados do arquivo e calcula a idade do aluno na data da avaliação."""
    df = pd.read_excel(uploaded_file, engine="openpyxl")
    df.columns = df.columns.str.strip()

    df.rename(columns={
        "Nome completo do aluno:": "Aluno",
        "Unidade escolar de origem do encaminhamento": "Unidade",
        "Data da avaliação:": "Data_Avaliacao",
        "Data de Nascimento:": "Data_Nascimento",
        "Nome do professor e demais profissionais que responderam o formulário:": "Professor"
    }, inplace=True)

    if "Data_Nascimento" not in df or "Data_Avaliacao" not in df:
        return None

    df["Data_Nascimento"] = pd.to_datetime(df["Data_Nascimento"], errors="coerce")
    df["Data_Avaliacao"] = pd.to_datetime(df["Data_Avaliacao"], errors="coerce")

    df["Ano"], df["Meses"], df["Meses_Totais"] = zip(*df.apply(
        lambda row: calcular_idade(row["Data_Nascimento"], row["Data_Avaliacao"]) if pd.notna(row["Data_Nascimento"]) and pd.notna(row["Data_Avaliacao"]) else (None, None, None),
        axis=1
    ))

    return df

def calcular_idade(data_nascimento, data_avaliacao):
    """Calcula a idade do aluno na data da avaliação."""
    if pd.isnull(data_nascimento) or pd.isnull(data_avaliacao):
        return None, None, None

    idade_anos = data_avaliacao.year - data_nascimento.year
    idade_meses = data_avaliacao.month - data_nascimento.month

    if data_avaliacao.day < data_nascimento.day:
        idade_meses -= 1

    if idade_meses < 0:
        idade_anos -= 1
        idade_meses += 12

    meses_totais = (idade_anos * 12) + idade_meses

    # Ajuste de arredondamento de até 2 dias
    diferenca_dias = (data_avaliacao - data_nascimento).days
    if 0 < diferenca_dias % 30 <= 2:
        meses_totais += 1
        idade_meses += 1

    return idade_anos, idade_meses, meses_totais

CORES_FIXAS_STATUS = {
    "Sem atraso ✅": "#2E7D32",
    "Alerta para atraso ⚠️": "#FFC107",
    "Possível Déficit 🚨": "#D32F2F"
}

CATEGORIAS_VALIDAS = ["Socialização", "Linguagem", "Cognição", "Auto cuidado", "Desenvolvimento Motor"]

def calcular_status_aluno(df, categoria, meses_faixa_etaria, pontuacao_esperada_manual=None):
    categorias = CATEGORIAS_VALIDAS if categoria == "Todas" else [categoria]
    df_resultado = []

    for _, row in df.iterrows():
        for categoria in categorias:
            colunas_categoria = [col for col in df.columns if col.startswith(categoria)]
            total_perguntas = len(colunas_categoria)

            if total_perguntas == 0:
                continue  

            pontos_obtidos = sum([
                1 if row[col] == "Sim" else 0.5 if row[col] == "Às vezes" else 0
                for col in colunas_categoria
            ])

            idade_meses = row["Meses"]
            if pd.isna(idade_meses):
                idade_meses = 0  

            if pd.isna(meses_faixa_etaria) or meses_faixa_etaria == 0:
                pontuacao_esperada = 0  
            else:
                pontuacao_esperada = (idade_meses * total_perguntas) / meses_faixa_etaria

            if pontuacao_esperada_manual is not None:
                pontuacao_esperada = pontuacao_esperada_manual

            if pontos_obtidos >= pontuacao_esperada:
                status = "Sem atraso ✅"
            elif pontos_obtidos >= (pontuacao_esperada - 2):
                status = "Alerta para atraso ⚠️"
            else:
                status = "Possível Déficit 🚨"

            df_resultado.append({
                "Aluno": row["Aluno"],
                "Categoria": categoria,
                "Pontuação Obtida": pontos_obtidos,
                "Pontuação Esperada": round(pontuacao_esperada, 2),
                "Status": status
            })

    df_resultado = pd.DataFrame(df_resultado)
    return df_resultado if not df_resultado.empty else None

def gerar_texto_analise(status_alunos):
    """Gera um parágrafo dinâmico com base nos resultados do aluno no inventário."""
    
    if status_alunos.empty:
        return "Não há dados suficientes para análise."

    categorias_por_status = {
        "Sem atraso ✅": [],
        "Alerta para atraso ⚠️": [],
        "Possível Déficit 🚨": []
    }

    for _, row in status_alunos.iterrows():
        categorias_por_status[row["Status"].strip()].append(row["Categoria"].strip())

    frases_status = {
        "Sem atraso ✅": "dentro do esperado para a faixa etária, indicando progresso adequado e consolidado",
        "Alerta para atraso ⚠️": "com indícios de dificuldades, apontando a necessidade de maior estimulação para fortalecer o desenvolvimento esperado",
        "Possível Déficit 🚨": "abaixo do esperado para a faixa etária, sugerindo possíveis dificuldades que demandam atenção e estimulação específica"
    }

    partes_texto = []

    for status, categorias in categorias_por_status.items():
        if categorias:
            lista_categorias = ", ".join(categorias[:-1]) + (" e " + categorias[-1] if len(categorias) > 1 else categorias[0])
            partes_texto.append(f"a(s) área(s) de {lista_categorias} está(ão) {frases_status[status]}")

    if partes_texto:
        return f"Com base nas respostas ao inventário, observa-se que {', '.join(partes_texto)}."
    
    return "Não há dados relevantes para análise."

def gerar_pdf(filtros, status_alunos, img_grafico,):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=20)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("Centro Municipal de Atendimento Especializado – CMAE", styles["Title"]))
    elements.append(Spacer(1, 12))

    for chave, valor in filtros.items():
        elements.append(Paragraph(f"<b>{chave}:</b> {valor}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    if not status_alunos.empty:
        elements.append(Paragraph("📊 Estatística(s) do(s) Aluno(s) - INVENTÁRIO PORTAGE", styles["Heading2"]))
        
        for aluno in status_alunos["Aluno"].unique():
            elements.append(Spacer(1, 12))
            dados_tabela = status_alunos[status_alunos["Aluno"] == aluno][["Categoria", "Pontuação Obtida", "Pontuação Esperada", "Status"]]
            tabela = Table([dados_tabela.columns.tolist()] + dados_tabela.values.tolist(), colWidths=[120, 110, 110, 120])
            tabela.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]))

        elements.append(Paragraph("Observação: É importante considerar que os resultados referem-se apenas a faixa etária analisada, sendo esta uma avaliação parcial, com vistas qualitativas."))
        elements.append(Spacer(1, 12))
        elements.append(tabela)
        elements.append(Spacer(1, 12))

    if img_grafico:
        img_path = "grafico_temp.png"
        with open(img_path, "wb") as f:
            f.write(img_grafico.getvalue())

        elements.append(Image(img_path, width=300, height=200))
        texto_analise = gerar_texto_analise(status_alunos)
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(texto_analise, styles["Normal"]))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("O Inventário Portage Operacionalizado (IPO) vem sendo respondido pelos professores dos Centros de Educação Infantil, de maneira adaptada e parcial, como forma de levantar dados e acompanhar o desenvolvimento das crianças."
        " Para investigação mais aprofundada, sugere-se a aplicação do Inventário Dimensional de Avaliação do Desenvolvimento Infantil - IDADI.", styles["Normal"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer

# Função para gerar o relatório em Word
def gerar_word(filtros, status_alunos, img_grafico):
    buffer = io.BytesIO()
    doc = Document()

    doc.add_heading("Centro Municipal de Atendimento Especializado – CMAE", level=1)

    # Adicionar filtros ao documento
    for chave, valor in filtros.items():
        doc.add_paragraph(f"{chave}: {valor}")

    doc.add_heading("📊 Estatísticas do(s) Aluno(s) - INVENTÁRIO PORTAGE", level=2)

    alunos_unicos = status_alunos["Aluno"].unique()
    for aluno in alunos_unicos:
        # Criando tabela de pontuação
        tabela = status_alunos[status_alunos["Aluno"] == aluno][["Categoria", "Pontuação Obtida", "Pontuação Esperada", "Status"]]
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Categoria"
        hdr_cells[1].text = "Pontuação Obtida"
        hdr_cells[2].text = "Pontuação Esperada"
        hdr_cells[3].text = "Status"

        for _, row in tabela.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row["Categoria"]
            row_cells[1].text = str(row["Pontuação Obtida"])
            row_cells[2].text = str(row["Pontuação Esperada"])
            row_cells[3].text = row["Status"]

    # Adiciona a análise do aluno
    texto_analise = gerar_texto_analise(status_alunos)
    doc.add_paragraph("\nAnálise Geral:")
    doc.add_paragraph(texto_analise)

    # Adicionar imagem do gráfico se disponível
    if img_grafico:
        img_path = "grafico_temp.png"
        with open(img_path, "wb") as f:
            f.write(img_grafico.getvalue())

        doc.add_paragraph("\n📊 Gráfico de Resultados:")
        doc.add_picture(img_path, width=Inches(4))

    doc.add_paragraph("\nO Inventário Portage Operacionalizado (IPO) vem sendo respondido pelos professores dos Centros de Educação Infantil, de maneira adaptada e parcial, como forma de levantar dados e acompanhar o desenvolvimento das crianças. Para investigação mais aprofundada, sugere-se a aplicação do Inventário Dimensional de Avaliação do Desenvolvimento Infantil - IDADI.")

    doc.save(buffer)
    buffer.seek(0)
    return buffer

def run_cmae_mode():
    st.title("📊 Painel Interativo de Avaliação (Modo CMAE)")

    uploaded_file = st.file_uploader("📂 Envie a planilha de respostas", type=["xlsx"])
    
    if not uploaded_file:
        st.info("🔍 Envie a planilha para iniciar a análise.")
        return
    
    df = carregar_dados(uploaded_file)
    if df is None:
        st.error("Erro ao carregar os dados. Verifique se as colunas necessárias estão presentes.")
        return
    
    df["Unidade"] = df["Unidade"].astype(str).str.strip()

    st.sidebar.header("🎯 **Filtros**")
    unidades = ["Todas"] + sorted(df["Unidade"].dropna().unique().tolist())
    unidade_selecionada = st.sidebar.selectbox("🏫 **Unidade Escolar**", unidades)

    categorias = ["Todas"] + CATEGORIAS_VALIDAS
    categoria_selecionada = st.sidebar.selectbox("🧩 **Categoria**", categorias)

    aluno_lista = df["Aluno"].dropna().unique().tolist()
    aluno_lista.insert(0, "Todos")
    aluno_selecionado = st.sidebar.selectbox("👦 **Pesquise um aluno**", aluno_lista)

    tipo_grafico = st.sidebar.selectbox(
    "📊 **Escolha o tipo de gráfico**",
    ["Barras", "Barras Empilhadas", "Barras por Resposta", "Pizza", "Linha"]
    )

    largura = st.sidebar.slider("📏 **Largura do Gráfico**", min_value=4, max_value=12, value=6, step=1)
    altura = st.sidebar.slider("📐 **Altura do Gráfico**", min_value=3, max_value=10, value=4, step=1)

    if st.sidebar.button("🔄 **Atualizar Dados**"):
        st.rerun()

    if unidade_selecionada != "Todas":
        df = df[df["Unidade"] == unidade_selecionada]
    if aluno_selecionado != "Todos":
        df = df[df["Aluno"] == aluno_selecionado]
    if not df.empty:
        aluno_info = df.iloc[0]
        filtros = {
            "Nome do Aluno": aluno_info.get("Aluno", "N/A"),
            "Unidade Escolar": aluno_info.get("Unidade", "N/A"),
            "Data da Avaliação": aluno_info.get("Data_Avaliacao", "N/A"),
            "Data de Nascimento": aluno_info.get("Data_Nascimento", "N/A"),
            "Idade": f"{aluno_info.get('Ano', 'N/A')} anos e {aluno_info.get('Meses', 'N/A')} meses",
            "Professor": aluno_info.get("Professor", "N/A")
        }

        st.write("### 📄 Informações do Aluno")
        for key, value in filtros.items():
            st.write(f"**{key}:** {value}")

    else:
        st.warning("Nenhum dado encontrado com os filtros selecionados.")
        return

    status_alunos = calcular_status_aluno(df, categoria_selecionada, 12)

    if status_alunos is not None and not status_alunos.empty:
        st.write(f"### 📊 Estatísticas para {aluno_selecionado} na Categoria: {categoria_selecionada}")
        st.dataframe(status_alunos)

        status_counts = status_alunos["Status"].value_counts()
        fig, ax = plt.subplots(figsize=(largura, altura))
        cores = [CORES_FIXAS_STATUS.get(status, "#000000") for status in status_counts.index]

        if tipo_grafico == "Barras":
            status_counts = status_alunos["Status"].value_counts()
            cores = [CORES_FIXAS_STATUS.get(status, "#000000") for status in status_counts.index]
            ax.bar(status_counts.index, status_counts.values, color=cores)
            ax.set_ylabel("Quantidade")
            ax.set_title("Distribuição dos Status")

        elif tipo_grafico == "Barras Empilhadas":
            df_stack = status_alunos.groupby(["Categoria", "Status"]).size().unstack(fill_value=0)
            cores_stack = [CORES_FIXAS_STATUS.get(status, "#000000") for status in df_stack.columns]
            df_stack.plot(kind="bar", stacked=True, ax=ax, color=cores_stack)
            ax.set_ylabel("Quantidade")
            ax.set_title("Distribuição de Status por Categoria")
            ax.legend(title="Status", bbox_to_anchor=(1.05, 1), loc="upper left")
            plt.tight_layout()

        elif tipo_grafico == "Pizza":
            status_counts = status_alunos["Status"].value_counts()
            cores = [CORES_FIXAS_STATUS.get(status, "#000000") for status in status_counts.index]
            ax.pie(status_counts.values, labels=status_counts.index, autopct='%1.1f%%', colors=cores)
            ax.axis("equal")
            ax.set_title("Distribuição dos Status")

        elif tipo_grafico == "Linha":
            status_counts = status_alunos["Status"].value_counts()
            ax.plot(status_counts.index, status_counts.values, marker='o', linestyle='-')
            ax.set_ylabel("Quantidade")
            ax.set_title("Distribuição dos Status")
        
        elif tipo_grafico == "Barras por Resposta":
            categorias_ativas = CATEGORIAS_VALIDAS if categoria_selecionada == "Todas" else [categoria_selecionada]

            # Inicializa um dicionário para contar "Sim", "Às vezes" e "Não" por categoria
            dados_respostas = {cat: {"Sim": 0, "Às vezes": 0, "Não": 0} for cat in categorias_ativas}

            for _, row in df.iterrows():
                for cat in categorias_ativas:
                    colunas = [col for col in df.columns if col.startswith(cat)]
                    for col in colunas:
                        valor = row[col]
                        if valor in ["Sim", "Às vezes", "Não"]:
                            dados_respostas[cat][valor] += 1

            # Converte para DataFrame para facilitar a plotagem
            df_resposta = pd.DataFrame(dados_respostas).T[["Sim", "Às vezes", "Não"]]

            cores_resp = ["#2E7D32", "#FFC107", "#D32F2F"]
            df_resposta.plot(kind="bar", stacked=True, ax=ax, color=cores_resp)

            ax.set_ylabel("Quantidade de Respostas")
            ax.set_title("Distribuição das Respostas por Categoria")
            ax.legend(title="Resposta", bbox_to_anchor=(1.05, 1), loc="upper left")
            plt.tight_layout()

        st.pyplot(fig)

        buffer_grafico = io.BytesIO()
        fig.savefig(buffer_grafico, format="png")
        buffer_grafico.seek(0)

        st.download_button(
            "📥 Baixar Relatório Completo (PDF)",
            gerar_pdf(filtros, status_alunos, buffer_grafico),
            file_name="relatorio_CMAE.pdf",
            mime="application/pdf"
        )
        
        # Criar botão de download apenas para o gráfico
        st.download_button(
            "📥 Baixar Gráfico",
            buffer_grafico,
            file_name="grafico_CMAE.png",
            mime="image/png"
        )
        buffer_word = gerar_word(filtros, status_alunos, buffer_grafico)
        st.download_button(
        "📥 Baixar Relatório Completo (Word)",
        buffer_word,
        file_name="relatorio_CMAE.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )