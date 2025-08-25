# sme_mode_v2.py
# ------------------------------------------------------------
# Streamlit – Painel Interativo (Modo SME) - V2
# - Aceita XLSX/CSV, escolha de aba
# - Detecta perguntas 1-5 (inteiro, "5 - Excelente", "⭐⭐⭐⭐", etc.)
# - Lida com colunas abertas (texto longo)
# - Agrupamento por Curso/Modalidade
# - Relatório .docx com gráficos (seek(0) fix) e Inches para largura
# ------------------------------------------------------------

import io
import re
import os
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# tenta importar quebrar_rotulo do seu utils; se não houver, usa fallback
try:
    from utils import quebrar_rotulo  # deve receber string e devolver string com quebras
except Exception:
    def quebrar_rotulo(s, largura=18):
        """Fallback simples para quebrar labels longos em múltiplas linhas."""
        if not isinstance(s, str):
            s = str(s)
        palavras, linhas, atual = s.split(), [], ""
        for w in palavras:
            if len(atual) + len(w) + 1 <= largura:
                atual = (atual + " " + w).strip()
            else:
                linhas.append(atual)
                atual = w
        if atual:
            linhas.append(atual)
        return "\n".join(linhas)

# -----------------------------
# Configurações visuais fixas
# -----------------------------
CORES_FIXAS = {
    # escalas antigas (se aparecerem)
    "Concordo plenamente": "#2E7D32", "Concordo": "#66BB6A", "Não sei": "#FFEB3B",
    "Discordo": "#FF7043", "Discordo plenamente": "#D32F2F",
    "Muito Satisfeito(a)": "#2E7D32", "Satisfeito(a)": "#66BB6A",
    "Regular": "#FFEB3B", "Insatisfeito(a)": "#FF7043",
}

# para classificação 1-5
PALETA_RATING = {
    1: "#D32F2F", 2: "#FF7043", 3: "#FFEB3B", 4: "#66BB6A", 5: "#2E7D32"
}

# colunas que não são perguntas (identificação/metadados mais comuns)
COLUNAS_META_CANDIDATAS = {
    "carimbo de data/hora", "timestamp", "nome", "nome completo",
    "curso", "modalidade (maior carga horária)", "modalidade",
    "email", "e-mail", "turma", "setor"
}

# padrões que ajudam a identificar perguntas (texto longo com ?)
PADRAO_PERGUNTA = re.compile(r"\?$", re.IGNORECASE)

# coluna aberta típica (ajuste os gatilhos conforme seus formulários)
PADROES_ABERTAS = [
    "pontos positivos", "pontos que poderiam ser melhorados",
    "observações", "comentários", "sugestões"
]


# -----------------------------
# Parsing de notas (1–5)
# -----------------------------
def parse_rating_cell(x):
    """
    Converte diferentes formatos para inteiros 1–5.
    Exemplos aceitos:
      - 1, 2, 3, 4, 5 (ints/strings)
      - "5 - Excelente", "3 – Regular"
      - "⭐⭐", "⭐⭐⭐⭐⭐"
      - "4.0"
    Retorna None se não identificar uma nota válida.
    """
    if pd.isna(x):
        return None

    # numérico direto
    if isinstance(x, (int, float)):
        xi = int(round(x))
        return xi if 1 <= xi <= 5 else None

    s = str(x).strip()

    # "5 - Excelente" → pega o primeiro número
    m = re.match(r"^\s*([1-5])\b", s)
    if m:
        return int(m.group(1))

    # "⭐⭐⭐" → conta estrelas
    if "⭐" in s:
        count = s.count("⭐")
        return count if 1 <= count <= 5 else None

    # "4.0"
    try:
        xf = float(s.replace(",", "."))
        xi = int(round(xf))
        if 1 <= xi <= 5:
            return xi
    except Exception:
        pass

    # mapeamentos livres (se quiser, amplie)
    MAP = {
        "péssimo": 1, "ruim": 2, "regular": 3,
        "bom": 4, "excelente": 5
    }
    s_low = s.lower()
    for k, v in MAP.items():
        if k in s_low:
            return v

    return None


def coluna_parece_aberta(nome_coluna):
    """Heurística para detectar perguntas abertas."""
    s = (nome_coluna or "").lower()
    if any(p in s for p in PADROES_ABERTAS):
        return True
    # se termina com interrogação e costuma gerar textos longos
    return False


def detectar_colunas(df):
    """
    Retorna:
      - meta_cols: colunas de identificação (curso, modalidade, etc.)
      - perguntas_numericas: perguntas que aceitam nota 1-5 (ou que podem ser parseadas)
      - perguntas_abertas: perguntas abertas (texto)
    """
    cols = list(df.columns)
    # meta: nomes exatos ou aproximações
    meta_cols = []
    for c in cols:
        c_low = str(c).strip().lower()
        if c_low in COLUNAS_META_CANDIDATAS:
            meta_cols.append(c)
        elif c_low.startswith("modalidade"):
            meta_cols.append(c)
        elif c_low in {"curso", "nome", "nome completo"}:
            meta_cols.append(c)
        elif c_low in {"carimbo de data/hora", "timestamp"}:
            meta_cols.append(c)

    # perguntas candidatas = restantes
    restantes = [c for c in cols if c not in meta_cols]

    perguntas_numericas, perguntas_abertas = [], []
    for c in restantes:
        serie = df[c]

        # tenta parsear como 1-5
        amostra = serie.dropna().head(20)
        parsed = [parse_rating_cell(v) for v in amostra]
        # se metade ou mais da amostra vira 1-5, consideramos numérica (estrelas)
        if len(parsed) > 0 and sum(v is not None for v in parsed) >= max(3, len(parsed)//2):
            perguntas_numericas.append(c)
            continue

        # se for texto e parecer "aberta"
        if serie.dtype == "object" and (coluna_parece_aberta(c) or PADRAO_PERGUNTA.search(str(c).strip()) is not None):
            perguntas_abertas.append(c)
            continue

        # fallback: se for object mas com poucas categorias e não meta, ainda pode ser pergunta categórica (antigo Likert)
        # mantemos fora nesta versão para foco no 1–5. (Se quiser, pode reincluir aqui.)

    return meta_cols, perguntas_numericas, perguntas_abertas


def preparar_dataframe(df):
    """Padroniza cabeçalhos e retorna df com colunas mínimas renomeadas onde possível."""
    # normaliza headers
    df.columns = df.columns.str.strip()

    # tenta renomear algumas chaves para padrão
    ren = {}
    for c in df.columns:
        cl = c.strip().lower()
        if cl == "selecione o curso" or cl == "curso":
            ren[c] = "Curso"
        elif cl.startswith("modalidade"):
            ren[c] = "Modalidade"
        elif cl in {"carimbo de data/hora", "timestamp"}:
            ren[c] = "Carimbo de data/hora"
        elif cl in {"nome", "nome completo"}:
            ren[c] = "Nome Completo"
    if ren:
        df = df.rename(columns=ren)

    return df


def grafico_barras_contagem(index_labels, valores, ax, titulo=None, paleta=None, anotar=True):
    """Barras simples com rótulos quebrados e contagem anotada."""
    pal = None
    if paleta:
        pal = [paleta.get(lbl, "#999999") for lbl in index_labels]

    sns.barplot(x=index_labels, y=valores, palette=pal, ax=ax)
    ax.set_xlabel("")
    ax.set_ylabel("Quantidade")
    ax.set_xticklabels([quebrar_rotulo(str(x)) for x in index_labels], ha="center")
    if anotar:
        for p in ax.patches:
            ax.annotate(
                f"{int(p.get_height())}",
                (p.get_x() + p.get_width() / 2, p.get_height()),
                ha="center", va="center", size=10, xytext=(0, 8), textcoords="offset points"
            )
    if titulo:
        ax.set_title(titulo)


def inserir_figura_no_doc(doc, fig, titulo):
    doc.add_heading(str(titulo), level=2)
    img_buffer = io.BytesIO()
    fig.savefig(img_buffer, format='png', bbox_inches='tight')
    img_buffer.seek(0)  # IMPORTANTE
    doc.add_picture(img_buffer, width=Inches(6))  # ~80% da página A4


def run_sme_mode():
    st.title("📊 Painel Interativo de Avaliação de Cursos (Modo SME) — V2")
    st.markdown("""
        <style>
            .main { background-color: #f0f2f6; }
            h1 { color: #2c3e50; text-align: center; }
        </style>
    """, unsafe_allow_html=True)

    st.sidebar.header("📂 Dados")
    uploaded = st.sidebar.file_uploader("Envie a planilha de respostas (XLSX ou CSV)", type=["xlsx", "csv"])

    df = None
    sheet_name = None

    if uploaded:
        filename = uploaded.name.lower()
        if filename.endswith(".xlsx"):
            # permite escolher a aba
            try:
                xls = pd.ExcelFile(uploaded)
                sheet_name = st.sidebar.selectbox("Aba da planilha", xls.sheet_names, index=0)
                df = pd.read_excel(xls, sheet_name=sheet_name)
            except Exception as e:
                st.error(f"Erro ao ler XLSX: {e}")
                return
        else:
            # CSV
            try:
                # tenta detectar sep padrão; ajuste se necessário
                df = pd.read_csv(uploaded)
            except Exception:
                uploaded.seek(0)
                df = pd.read_csv(uploaded, sep=";")

        if df is None or df.empty:
            st.warning("Não foi possível carregar dados.")
            return

        df = preparar_dataframe(df)

        # mostra amostra
        with st.expander("🔎 Prévia dos dados (primeiras linhas)"):
            st.dataframe(df.head(10), use_container_width=True)

        # detectar colunas
        meta_cols, perguntas_numericas, perguntas_abertas = detectar_colunas(df)

        # filtros
        st.sidebar.header("🎛️ Filtros")
        curso_col = "Curso" if "Curso" in df.columns else None
        modalidade_col = "Modalidade" if "Modalidade" in df.columns else None

        if curso_col:
            cursos = ["Todos"] + sorted([c for c in df[curso_col].dropna().unique()])
            curso_filtro = st.sidebar.selectbox("Curso", cursos, index=0)
        else:
            curso_filtro = "Todos"

        if modalidade_col:
            modalidades = ["Todas"] + sorted([c for c in df[modalidade_col].dropna().unique()])
            modalidade_filtro = st.sidebar.selectbox("Modalidade", modalidades, index=0)
        else:
            modalidade_filtro = "Todas"

        if st.sidebar.button("🔄 Recarregar"):
            st.rerun()

        # aplica filtros
        df_filtrado = df.copy()
        if curso_col and curso_filtro != "Todos":
            df_filtrado = df_filtrado[df_filtrado[curso_col] == curso_filtro]
        if modalidade_col and modalidade_filtro != "Todas":
            df_filtrado = df_filtrado[df_filtrado[modalidade_col] == modalidade_filtro]

        st.sidebar.markdown("---")
        st.sidebar.write(f"📊 Total de respostas (filtro aplicado): **{len(df_filtrado)}**")

        # seleção de perguntas
        st.sidebar.header("📝 Perguntas")
        lista_perguntas = perguntas_numericas + perguntas_abertas
        if not lista_perguntas:
            st.info("Não foram detectadas perguntas. Verifique se sua planilha possui colunas de perguntas (1–5) ou abertas.")
            return

        perguntas_escolhidas = st.sidebar.multiselect(
            "Selecione as perguntas para análise",
            options=lista_perguntas,
            default=perguntas_numericas[:5] if perguntas_numericas else lista_perguntas[:3]
        )

        # tipo de gráfico para perguntas numéricas
        tipo_grafico = st.sidebar.selectbox("Gráfico para perguntas 1–5", ["Barras", "Pizza"], index=0)

        # agrupar por (opcional)
        st.sidebar.header("📚 Agrupamento")
        group_by = st.sidebar.multiselect(
            "Agrupar por",
            options=[c for c in ["Curso", "Modalidade"] if c in df_filtrado.columns],
            default=[]
        )

        # sliders de tamanho
        largura = st.sidebar.slider("Largura do gráfico (inches aprox.)", 4, 14, 7)
        altura = st.sidebar.slider("Altura do gráfico (inches aprox.)", 3, 8, 4)

        # DOCX
        doc = Document()
        titulo_doc = f"Relatório de Avaliação - {curso_filtro if curso_filtro!='Todos' else 'Todos os cursos'}"
        if modalidade_col and modalidade_filtro != "Todas":
            titulo_doc += f" | Modalidade: {modalidade_filtro}"
        doc.add_heading(titulo_doc, level=1)

        # ======= RESUMO GERAL (só notas) =======
        if perguntas_numericas:
            st.header("📈 Resumo Geral (médias 1–5)")
            resumo = {}
            for col in perguntas_numericas:
                ratings = df_filtrado[col].apply(parse_rating_cell)
                serie = pd.Series([r for r in ratings if r is not None])
                if len(serie) > 0:
                    resumo[col] = {
                        "Média": round(serie.mean(), 2),
                        "N": int(serie.count())
                    }
            if resumo:
                df_resumo = pd.DataFrame(resumo).T.sort_values("Média", ascending=False)
                st.dataframe(df_resumo, use_container_width=True)
                # índice agregado (média de médias ponderada por N)
                if not df_resumo.empty:
                    agg = (df_resumo["Média"] * df_resumo["N"]).sum() / max(df_resumo["N"].sum(), 1)
                    st.success(f"⭐ **Índice agregado (média geral ponderada)**: {agg:.2f} / 5")
            else:
                st.info("Não foi possível calcular médias. Verifique as colunas numéricas.")

        # ======= LOOP DAS PERGUNTAS ESCOLHIDAS =======
        for col in perguntas_escolhidas:
            st.subheader(col)

            if col in perguntas_numericas:
                # distribuições 1–5
                serie = df_filtrado[col].apply(parse_rating_cell)
                serie = serie.dropna().astype(int)
                # garante presença de 1..5 mesmo sem respostas
                contagem = serie.value_counts().reindex([1,2,3,4,5], fill_value=0)

                if group_by:
                    # gráfico por grupo (barras empilhadas simples por categoria)
                    for g in group_by:
                        st.markdown(f"**Distribuição por {g}**")
                        fig, ax = plt.subplots(figsize=(largura, altura))
                        # tabela p/ cada categoria de g
                        tmp = df_filtrado[[g, col]].copy()
                        tmp["rating"] = tmp[col].apply(parse_rating_cell)
                        tmp = tmp.dropna(subset=["rating"])
                        tb = (tmp.pivot_table(index=g, columns="rating", values=col, aggfunc="count", fill_value=0)
                                  .reindex(columns=[1,2,3,4,5], fill_value=0))
                        tb.plot(kind="bar", stacked=True, ax=ax, color=[PALETA_RATING[k] for k in [1,2,3,4,5]])
                        ax.set_xlabel("")
                        ax.set_ylabel("Quantidade")
                        ax.set_xticklabels([quebrar_rotulo(str(x)) for x in tb.index], ha="center")
                        ax.legend(title="Nota", bbox_to_anchor=(1.01, 1), loc="upper left")
                        st.pyplot(fig, use_container_width=True)
                        inserir_figura_no_doc(doc, fig, f"{col} — Distribuição por {g}")

                # gráfico principal da pergunta
                fig, ax = plt.subplots(figsize=(largura, altura))
                if tipo_grafico == "Barras":
                    grafico_barras_contagem(contagem.index.tolist(), contagem.values.tolist(), ax,
                                            titulo="Distribuição de notas", paleta=PALETA_RATING)
                else:
                    # Pizza
                    vals = contagem.values.tolist()
                    labels = contagem.index.tolist()
                    colors = [PALETA_RATING.get(k, "#999999") for k in labels]
                    ax.pie(vals, labels=labels, autopct='%1.1f%%', startangle=140, colors=colors)
                    ax.set_title("Distribuição de notas")
                st.pyplot(fig, use_container_width=True)
                inserir_figura_no_doc(doc, fig, f"{col} — Distribuição geral")

                # média por grupo (se houver)
                if group_by:
                    fig, ax = plt.subplots(figsize=(largura, altura))
                    tmp = df_filtrado[group_by + [col]].copy()
                    for g in group_by:
                        tmp[g] = tmp[g].astype(str)
                    tmp["rating"] = tmp[col].apply(parse_rating_cell)
                    tmp = tmp.dropna(subset=["rating"])
                    # média por combinação de grupos selecionados
                    med = tmp.groupby(group_by)["rating"].mean().sort_values(ascending=False)
                    sns.barplot(x=med.index.astype(str), y=med.values, ax=ax)
                    ax.set_xlabel("")
                    ax.set_ylabel("Média (1–5)")
                    ax.set_xticklabels([quebrar_rotulo(str(x)) for x in med.index], ha="center")
                    for p in ax.patches:
                        ax.annotate(
                            f"{p.get_height():.2f}",
                            (p.get_x() + p.get_width()/2, p.get_height()),
                            ha="center", va="center", size=10, xytext=(0, 8), textcoords="offset points"
                        )
                    ax.set_title("Média por grupo")
                    st.pyplot(fig, use_container_width=True)
                    inserir_figura_no_doc(doc, fig, f"{col} — Média por grupo")

            elif col in perguntas_abertas:
                # texto livre: mostra amostra e top palavras (contagem simples)
                serie_txt = df_filtrado[col].dropna().astype(str)
                st.write(f"Respostas (N={len(serie_txt)}):")
                st.dataframe(serie_txt.to_frame(col), use_container_width=True)

                # contagem simples de palavras (exclui muito curtas)
                todas = " ".join(serie_txt.tolist()).lower()
                # remove pontuação básica
                todas = re.sub(r"[.,;:!?/()\"'“”·\-—–_]", " ", todas)
                tokens = [t for t in todas.split() if len(t) >= 4]
                # stopwords básicas PT (enxuta; ajuste/importe NLTK se quiser algo melhor)
                stop = {
                    "para","como","isso","essa","esse","esta","este","aquele","aquela","aqui","ali","tudo",
                    "muito","também","pois","onde","quando","então","porque","pela","pelo","pela","pelos",
                    "sobre","entre","apesar","com","sem","mais","menos","cada","todos","todas","foram",
                    "curso","aulas","professor","instrutor","facilitador","palestrante"  # ajuste se quiser manter
                }
                tokens = [t for t in tokens if t not in stop]
                cont = pd.Series(tokens).value_counts().head(20)

                if not cont.empty:
                    fig, ax = plt.subplots(figsize=(largura, altura))
                    grafico_barras_contagem(cont.index.tolist(), cont.values.tolist(), ax, titulo="Palavras mais frequentes")
                    st.pyplot(fig, use_container_width=True)
                    inserir_figura_no_doc(doc, fig, f"{col} — Palavras mais frequentes")

        # ======= EXPORTAÇÃO DOCX =======
        st.sidebar.markdown("---")
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        st.sidebar.download_button(
            "📥 Baixar relatório (DOCX)",
            doc_buffer.getvalue(),
            file_name="Relatorio_Avaliacao.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        st.info("🔍 Envie uma planilha (.xlsx ou .csv) para iniciar a análise.")
