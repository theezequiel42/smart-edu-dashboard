import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import io
import tempfile
import os
import openpyxl

# Configuração da página
st.set_page_config(page_title="Análise de Cursos", layout="wide")

# Estilização da interface
st.markdown("""
    <style>
        .main {
            background-color: #f0f2f6;
        }
        h1 {
            color: #2c3e50;
            text-align: center;
        }
        .sidebar-button {
            background-color: #2E7D32 !important;
            color: white !important;
            font-size: 16px !important;
            font-weight: bold !important;
            text-align: center;
            width: 100%;
            padding: 10px;
            border-radius: 8px;
        }
    </style>
""", unsafe_allow_html=True)

# Dicionário de cores fixas para opções de resposta
cores_fixas = {
    "Concordo plenamente": "#2E7D32",  # Verde escuro
    "Concordo": "#66BB6A",  # Verde claro
    "Não sei": "#FFEB3B",  # Amarelo
    "Discordo": "#FF7043",  # Laranja
    "Discordo plenamente": "#D32F2F",  # Vermelho
    "Muito Satisfeito(a)": "#2E7D32",
    "Satisfeito(a)": "#66BB6A",
    "Regular": "#FFEB3B",
    "Insatisfeito(a)": "#FF7043",
    "Muito motivado(a)": "#2E7D32",
    "Motivado(a)": "#66BB6A",
    "Ansioso(a) ou inseguro(a)": "#FF7043",
    "Sem expectativas claras": "#D32F2F",
    "Interesse no tema do curso": "#1E88E5",  # Azul forte
    "Interesse nas horas para ampliação / Foi o que consegui me inscrever": "#42A5F5",  # Azul médio
    "Não tinha muito interesse mas dei uma oportunidade ao tema": "#90CAF9"  # Azul claro
}

# Título do painel
st.title("📊 Painel Interativo de Avaliação de Cursos")

# Upload do arquivo
uploaded_file = st.file_uploader("📂 Envie o arquivo da planilha de respostas", type=["xlsx"])

if uploaded_file:
    # Carregar os dados da planilha
    df = pd.read_excel(uploaded_file, sheet_name=None)  # Carregar todas as abas
    sheet_name = list(df.keys())[0]  # Pegar a primeira aba disponível
    df = df[sheet_name]

    # Padronizar nomes das colunas
    df.columns = df.columns.str.strip()
    df.columns = df.columns.str.lower()

    # Converter colunas com datas para strings
    for col in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[col]):
            df[col] = df[col].astype(str)

    # Exibir colunas disponíveis para depuração
    st.write("Colunas disponíveis na planilha:", df.columns.tolist())

    # Renomeando colunas para facilitar
    df.rename(columns={
        'selecione o curso': 'curso',
        'satisfação geral: como você avalia sua satisfação geral com o curso?': 'satisfação geral',
        'conteúdo do curso [o curso foi relevante para suas atividades profissionais ou interesses]': 'o curso foi relevante para suas atividades profissionais ou interesses',
        'habilidade e didática do instrutor [o instrutor foi um palestrante/demonstrador eficiente]': 'o instrutor foi um palestrante/demonstrador eficiente'
    }, inplace=True, errors='ignore')

    # Barra lateral para filtros
    st.sidebar.header("🎓 Filtros")
    cursos = df['curso'].unique()
    curso_selecionado = st.sidebar.selectbox("📌 Selecione o Curso", cursos)

    # Botão para recarregar dados
    if st.sidebar.button("🔄 Recarregar Dados"):
        st.rerun()

    # Exibir estatísticas gerais
    st.sidebar.header("📊 Estatísticas Gerais")
    numeric_columns = df.select_dtypes(include=['number'])
    if not numeric_columns.empty:
        st.sidebar.write("Número total de respostas:", len(df))
        st.sidebar.write("Média das respostas (numéricas):", numeric_columns.describe().mean())
    else:
        st.sidebar.write("Não há colunas numéricas para análise estatística.")

    # Filtro de perguntas
    perguntas = [col for col in df.columns if col not in ['curso']]
    perguntas_selecionadas = st.sidebar.multiselect("📋 Selecione as perguntas para visualizar", perguntas, default=perguntas if perguntas else [])

    # Modo de exibição
    tipo_grafico = st.sidebar.selectbox("📈 Escolha o tipo de gráfico", ["Barras", "Pizza", "Linha"])

    # Criar um documento Word
    doc = Document()
    doc.add_heading(f"Relatório de Avaliação - {curso_selecionado}", level=1)

    arquivos_temp = []

    if perguntas_selecionadas:
        for coluna in perguntas_selecionadas:
            st.subheader(f"📊 {coluna}")

            # Sliders individuais para tamanho dos gráficos
            largura_grafico = st.slider(f"📏 Largura do gráfico ({coluna})", min_value=3, max_value=12, value=6)
            altura_grafico = st.slider(f"📐 Altura do gráfico ({coluna})", min_value=2, max_value=10, value=4)
            
            # Aplicação das cores fixas
            contagem = df[coluna].value_counts()
            paleta = {x: cores_fixas.get(x, "#999999") for x in contagem.index}

            # Gerar gráfico com base no tipo selecionado
            fig, ax = plt.subplots(figsize=(largura_grafico, altura_grafico))
            if tipo_grafico == "Barras":
                sns.barplot(x=contagem.index, y=contagem.values, hue=contagem.index, palette=paleta, legend=False, ax=ax)
                for bar, label in zip(ax.patches, contagem.index):
                    ax.annotate(f'{bar.get_height()}', (bar.get_x() + bar.get_width() / 2, bar.get_height()),
                                ha='center', va='center', size=10, xytext=(0, 8), textcoords='offset points')
                ax.set_xlabel("")
                ax.set_ylabel("Quantidade")
                ax.set_title(coluna, fontsize=10, fontweight='bold')
            elif tipo_grafico == "Pizza":
                ax.pie(contagem.values, labels=contagem.index, colors=[paleta[x] for x in contagem.index], autopct='%1.1f%%', startangle=140)
            elif tipo_grafico == "Linha":
                ax.plot(contagem.index, contagem.values, marker='o', color='#1E88E5')
            st.pyplot(fig, use_container_width=True)

            # Criar um arquivo temporário para a imagem
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                fig.savefig(tmp_file.name, format='png', bbox_inches='tight')
                tmp_file_path = tmp_file.name
                arquivos_temp.append(tmp_file_path)

            # Adicionar gráfico ao documento
            doc.add_heading(coluna, level=2)
            doc.add_paragraph("Número de respostas para cada opção:")
            doc.add_picture(tmp_file_path, width=doc.sections[0].page_width * 0.8)

            # Botão para baixar gráfico individualmente
            with open(tmp_file_path, "rb") as f:
                st.download_button(f"📥 Baixar gráfico - {coluna}", f, file_name=f"{coluna}.png", mime="image/png")

    # Remover arquivos temporários após o uso
    for arq in arquivos_temp:
        if os.path.exists(arq):
            os.remove(arq)

    # Botão para exportar Word na barra lateral
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    st.sidebar.markdown("### 📥 Baixar Relatório")
    st.sidebar.download_button("📥 Baixar relatório em Word", data=doc_buffer.getvalue(), file_name=f"Relatorio_{curso_selecionado}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", help="Clique para baixar o relatório gerado.", key="download_word", use_container_width=True)

    # Botão para exportar os dados filtrados para Excel
    st.sidebar.markdown("### 📤 Exportar Dados")
    excel_buffer = io.BytesIO()
    df_filtrado = df[df['curso'] == curso_selecionado]
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df_filtrado.to_excel(writer, index=False, sheet_name="Avaliações")
    excel_buffer.seek(0)
    st.sidebar.download_button("📥 Baixar dados filtrados em Excel", data=excel_buffer.getvalue(), file_name="dados_filtrados.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
else:
    st.info("🔍 Por favor, envie a planilha para iniciar a análise.")
